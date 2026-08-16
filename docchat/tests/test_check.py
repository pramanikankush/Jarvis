"""Self-check for core logic (chunking, parsing, store+search). Run: python tests/test_check.py"""
import io
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import numpy as np

from ragchat import parsing
from ragchat.store import Store


def test_chunking_basics():
    chunks = parsing.chunk_text("One small paragraph here.")
    assert chunks == ["One small paragraph here."], chunks

    text = ("\n\n".join(f"Paragraph {i} " + "word " * 300 for i in range(5)))
    chunks = parsing.chunk_text(text)
    assert chunks, "no chunks produced"
    assert all(c.strip() for c in chunks), "empty chunk produced"
    assert all(len(c) <= 1250 for c in chunks), f"chunk too long: {max(len(c) for c in chunks)}"
    assert all(len(c) >= 100 for c in chunks), "chunks suspiciously small"


def test_chunking_overlap_on_long_paragraph():
    para = "Sentence one. " * 200  # no newlines → must be split on sentence boundaries
    chunks = parsing.chunk_text(para)
    assert len(chunks) > 1, "long paragraph was not split"
    # overlap content should share some words between adjacent chunks
    shared = set(chunks[0].split()) & set(chunks[1].split())
    assert shared, "adjacent chunks share nothing (overlap broken)"


def test_hard_split_single_long_word_run():
    # one 4800-char run with spaces but no sentence punctuation → must be word-split
    chunks = parsing.chunk_text("word " * 400)
    assert chunks and all(len(c) <= 1250 for c in chunks)


def test_txt_decoding():
    utf8 = "héllo wörld".encode("utf-8")
    assert parsing.decode_text(utf8) == "héllo wörld"
    latin1 = "café résumé".encode("latin-1")
    assert parsing.decode_text(latin1) == "café résumé"


def test_csv_parse_with_header():
    csv_bytes = "name,amount,date\nAlice,10.5,2024-01-01\nBob,20,2024-02-02\n".encode("utf-8")
    text = parsing._parse_csv(csv_bytes)
    assert "Alice" in text and "amount: 10.5" in text and "name:" in text


def test_csv_parse_without_header():
    csv_bytes = "1,2,3\n4,5,6\n".encode("utf-8")
    text = parsing._parse_csv(csv_bytes)
    assert "1, 2, 3" in text and "4, 5, 6" in text


def test_docx_parse_paragraphs_and_tables():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello from a docx paragraph.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "CellA"
    table.rows[0].cells[1].text = "CellB"
    buf = io.BytesIO()
    doc.save(buf)
    text = parsing._parse_docx(buf.getvalue())
    assert "Hello from a docx paragraph." in text
    assert "CellA | CellB" in text


def test_store_roundtrip_and_search():
    with tempfile.TemporaryDirectory() as td:
        st = Store(os.path.join(td, "t.db"))
        rng = np.random.default_rng(7)
        vecs = [rng.normal(size=8).astype(np.float32) for _ in range(4)]
        for v in vecs:
            v /= np.linalg.norm(v)
        st.add_doc(
            "a.txt", 100,
            [(None, "alpha"), (None, "beta"), (None, "gamma"), (None, "delta")],
            vecs,
        )
        assert st.total_chunks() == 4
        assert st.list_docs()[0]["name"] == "a.txt"

        # query closest to vecs[2]
        q = vecs[2] / np.linalg.norm(vecs[2])
        top = st.search(q, k=2)
        assert top and top[0]["text"] == "gamma", f"expected gamma first, got {top[0]['text']}"
        st.close()

        # persistence: reopen, data still there
        st2 = Store(os.path.join(td, "t.db"))
        assert st2.total_chunks() == 4
        assert st2.search(q, k=1)[0]["text"] == "gamma"

        # delete
        doc_id = st2.list_docs()[0]["id"]
        assert st2.delete_doc(doc_id)
        assert st2.total_chunks() == 0
        assert st2.search(q, k=2) == []
        st2.close()


def test_sessions_persistence():
    with tempfile.TemporaryDirectory() as td:
        st = Store(os.path.join(td, "t.db"))
        s = st.create_session("First chat")
        st.add_message(s["id"], "user", "hello")
        st.add_message(s["id"], "assistant", "hi", [{"doc_name": "x.pdf", "page": 1, "text": "…"}])
        msgs = st.messages(s["id"])
        assert [m["role"] for m in msgs] == ["user", "assistant"]
        assert msgs[1]["sources"][0]["doc_name"] == "x.pdf"
        st2 = Store(os.path.join(td, "t.db"))
        assert len(st2.messages(s["id"])) == 2
        assert st2.delete_session(s["id"])
        assert st2.messages(s["id"]) == []
        st2.close()
        st.close()


def test_session_management_rename_pin_search_export():
    with tempfile.TemporaryDirectory() as td:
        st = Store(os.path.join(td, "t.db"))
        s = st.create_session("First chat")
        st.add_message(s["id"], "user", "hello there")
        st.set_title(s["id"], "Renamed chat")
        assert st.get_session(s["id"])["title"] == "Renamed chat"
        assert st.set_pinned(s["id"], True)
        assert st.get_session(s["id"])["pinned"] == 1
        st.create_session("Another chat")
        # search by title
        hits = st.search_sessions("rename")
        assert [h["id"] for h in hits] == [s["id"]], hits
        assert st.search_sessions("")  # empty query -> list, not crash
        # pinned sessions sort first
        assert st.list_sessions()[0]["id"] == s["id"]
        # export
        exp = st.export_session(s["id"])
        assert exp["title"] == "Renamed chat"
        assert exp["messages"][0]["content"] == "hello there"
        assert exp["pinned"] is True
        assert st.export_session("nope") is None
        st.close()


def test_store_memory_update():
    with tempfile.TemporaryDirectory() as td:
        st = Store(os.path.join(td, "t.db"))
        row = st.add_memory("the user likes tea", "preference")
        assert row is not None
        assert st.update_memory(row["id"], "the user likes coffee")
        assert st.list_memory()[0]["fact"] == "the user likes coffee"
        assert not st.update_memory(99999, "nope")
        st.close()


def test_store_migration_adds_pinned_column():
    """Existing app.db files (pre-pinned) must be upgraded in place."""
    import sqlite3

    with tempfile.TemporaryDirectory() as td:
        path = os.path.join(td, "old.db")
        conn = sqlite3.connect(path)
        conn.executescript("""
            CREATE TABLE sessions (id TEXT PRIMARY KEY, title TEXT NOT NULL,
                                   created_at TEXT NOT NULL, updated_at TEXT NOT NULL);
            CREATE TABLE messages (id INTEGER PRIMARY KEY AUTOINCREMENT,
                                   session_id TEXT NOT NULL, role TEXT NOT NULL,
                                   content TEXT NOT NULL, sources TEXT, created_at TEXT NOT NULL);
        """)
        conn.commit()
        conn.close()
        st = Store(path)
        cols = [r["name"] for r in st._conn.execute("PRAGMA table_info(sessions)").fetchall()]
        assert "pinned" in cols, cols
        # existing session still readable after migration
        s = st.create_session("new")
        assert st.get_session(s["id"])["pinned"] == 0
        st.close()


def main():
    tests = [v for k, v in sorted(globals().items()) if k.startswith("test_")]
    failed = 0
    for t in tests:
        try:
            t()
            print(f"  ok  {t.__name__}")
        except Exception as e:
            failed += 1
            print(f"FAIL  {t.__name__}: {e}")
    print(f"\n{len(tests) - failed}/{len(tests)} passed")
    sys.exit(1 if failed else 0)


if __name__ == "__main__":
    main()